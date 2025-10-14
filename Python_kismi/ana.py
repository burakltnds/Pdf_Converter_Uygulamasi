import fitz
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def metni_cikar(pdf_path):
    """PDF'den tüm metni çıkarır."""
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n"
    doc.close()
    return text

def clean_text(text):
    """Metni temizler, cümlelere böler ve Word için formatlar."""
    text = re.sub(r'\s+', ' ', text).strip()
    lines = text.split()
    cleaned_lines = []

    for line in lines:
        line = line.strip()
        if not line:
            continue
        cleaned_lines.append(line)
    return cleaned_lines

def contains_unknown_characters(line):
    """Satırda tanınmayan veya bozuk karakter varsa True döndürür."""
    for ch in line:
        # Normal Latin, Türkçe veya temel noktalama karakteri değilse
        if not (32 <= ord(ch) <= 126 or ch in "çğıöşüÇĞİÖŞÜ.,:;!?()[]{}«»“”’‘-–—'\" "):
            return True
    return False

def save_to_word(lines, output_path):
    """Cümleleri Word'e ekler ve özel biçimlendirmeleri uygular."""
    doc = Document()
    paragraph = doc.add_paragraph()
    for line in lines:
        if line.startswith(("•", "", "▪", "-", "‣", "\uf0b7", "\u2022","1.","2.","3.","4.","5.")) or line.endswith(("i." , "I.","v.","V.")) or contains_unknown_characters(line) :
            # Madde işareti cümlesi
            paragraph = doc.add_paragraph(line)
            paragraph.style.font.name = "Calibri"
            paragraph.style.font.size = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)
        elif line.isupper():
            # Tamamı büyük harf cümle
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.style.font.name = "Calibri"
            paragraph.style.font.size = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.runs[0]
            run.bold = True
            paragraph = doc.add_paragraph()
        else:
            # Normal cümle
            paragraph.add_run(line + " ")
        
        
        paragraph.style.font.name = "Calibri"
        paragraph.style.font.size = Pt(10)
        

    doc.save(output_path)
    print(f"✅ Word dosyası oluşturuldu: {output_path}")